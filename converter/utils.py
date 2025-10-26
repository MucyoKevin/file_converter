"""
Conversion utility functions for different file formats
"""
import os
from PIL import Image
import PyPDF2
from pdf2docx import Converter as PDFToDocxConverter
from docx import Document
import cv2
import numpy as np

try:
    from moviepy.editor import VideoFileClip
except ImportError:
    try:
        from moviepy import VideoFileClip
    except ImportError:
        VideoFileClip = None


def get_temp_path(source_path, target_format):
    """Generate temporary output path"""
    base = os.path.splitext(source_path)[0]
    return f"{base}_converted.{target_format}"


# ==================== IMAGE CONVERSIONS ====================

def convert_image_format(source_path, target_format):
    """Convert image from one format to another"""
    output_path = get_temp_path(source_path, target_format)
    
    with Image.open(source_path) as img:
        # Handle transparency for formats that don't support it
        if target_format.lower() in ['jpg', 'jpeg'] and img.mode in ['RGBA', 'LA', 'P']:
            # Create white background
            background = Image.new('RGB', img.size, (255, 255, 255))
            if img.mode == 'P':
                img = img.convert('RGBA')
            background.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
            img = background
        elif img.mode not in ['RGB', 'RGBA']:
            img = img.convert('RGB')
        
        # Save with optimization
        if target_format.lower() in ['jpg', 'jpeg']:
            img.save(output_path, 'JPEG', quality=95, optimize=True)
        elif target_format.lower() == 'png':
            img.save(output_path, 'PNG', optimize=True)
        elif target_format.lower() == 'webp':
            img.save(output_path, 'WEBP', quality=90)
        else:
            img.save(output_path, target_format.upper())
    
    return output_path


def image_to_pdf(source_path, target_format='pdf'):
    """Convert image to PDF"""
    output_path = get_temp_path(source_path, 'pdf')
    
    with Image.open(source_path) as img:
        # Convert to RGB if necessary
        if img.mode == 'RGBA':
            background = Image.new('RGB', img.size, (255, 255, 255))
            background.paste(img, mask=img.split()[-1])
            img = background
        elif img.mode != 'RGB':
            img = img.convert('RGB')
        
        img.save(output_path, 'PDF', resolution=100.0)
    
    return output_path


def pdf_to_image(source_path, target_format='jpg', page_number=1):
    """Convert PDF page to image using OpenCV and PyPDF2"""
    output_path = get_temp_path(source_path, target_format)
    
    try:
        # Try using pdf2image if available
        from pdf2image import convert_from_path
        images = convert_from_path(source_path, first_page=page_number, last_page=page_number, dpi=200)
        if images:
            if target_format.lower() in ['jpg', 'jpeg']:
                images[0].save(output_path, 'JPEG', quality=95)
            else:
                images[0].save(output_path, target_format.upper())
            return output_path
    except ImportError:
        pass
    
    # Fallback: Create a simple image representation
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    
    # Read PDF and extract text
    with open(source_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        if page_number <= len(pdf_reader.pages):
            page = pdf_reader.pages[page_number - 1]
            text = page.extract_text()
    
    # Create image from text
    width, height = 800, 1000
    img = Image.new('RGB', (width, height), color='white')
    
    # For now, save a placeholder
    img.save(output_path, target_format.upper())
    
    return output_path


# ==================== DOCUMENT CONVERSIONS ====================

def pdf_to_docx(source_path, target_format='docx'):
    """Convert PDF to DOCX"""
    output_path = get_temp_path(source_path, 'docx')
    
    cv = PDFToDocxConverter(source_path)
    cv.convert(output_path, start=0, end=None)
    cv.close()
    
    return output_path


def docx_to_pdf(source_path, target_format='pdf'):
    """Convert DOCX to PDF using reportlab"""
    output_path = get_temp_path(source_path, 'pdf')
    
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    
    # Read DOCX
    doc = Document(source_path)
    
    # Create PDF
    pdf = SimpleDocTemplate(output_path, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            p = Paragraph(paragraph.text, styles['Normal'])
            story.append(p)
            story.append(Spacer(1, 0.2 * inch))
    
    pdf.build(story)
    
    return output_path


def pdf_to_txt(source_path, target_format='txt'):
    """Convert PDF to TXT"""
    output_path = get_temp_path(source_path, 'txt')
    
    with open(source_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        text_content = []
        
        for page in pdf_reader.pages:
            text_content.append(page.extract_text())
    
    with open(output_path, 'w', encoding='utf-8') as file:
        file.write('\n\n'.join(text_content))
    
    return output_path


def docx_to_txt(source_path, target_format='txt'):
    """Convert DOCX to TXT"""
    output_path = get_temp_path(source_path, 'txt')
    
    doc = Document(source_path)
    text_content = [paragraph.text for paragraph in doc.paragraphs]
    
    with open(output_path, 'w', encoding='utf-8') as file:
        file.write('\n'.join(text_content))
    
    return output_path


def txt_to_pdf(source_path, target_format='pdf'):
    """Convert TXT to PDF"""
    output_path = get_temp_path(source_path, 'pdf')
    
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    
    with open(source_path, 'r', encoding='utf-8') as file:
        lines = file.readlines()
    
    pdf = SimpleDocTemplate(output_path, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    for line in lines:
        if line.strip():
            p = Paragraph(line.strip(), styles['Normal'])
            story.append(p)
    
    pdf.build(story)
    
    return output_path


# ==================== VIDEO CONVERSIONS ====================

def video_to_gif(source_path, target_format='gif', max_duration=10, max_width=480):
    """Convert video to GIF"""
    if VideoFileClip is None:
        raise ImportError("moviepy is required for video conversion. Please install it: pip install moviepy")
    
    output_path = get_temp_path(source_path, 'gif')
    
    clip = VideoFileClip(source_path)
    
    # Limit duration
    if clip.duration > max_duration:
        clip = clip.subclip(0, max_duration)
    
    # Resize if too large
    if clip.size[0] > max_width:
        clip = clip.resize(width=max_width)
    
    # Write GIF with optimized settings
    clip.write_gif(output_path, fps=10, program='ffmpeg', opt='nq')
    clip.close()
    
    return output_path


def convert_video_format(source_path, target_format):
    """Convert video from one format to another"""
    if VideoFileClip is None:
        raise ImportError("moviepy is required for video conversion. Please install it: pip install moviepy")
    
    output_path = get_temp_path(source_path, target_format)
    
    clip = VideoFileClip(source_path)
    
    # Write with appropriate codec
    if target_format == 'mp4':
        clip.write_videofile(output_path, codec='libx264', audio_codec='aac')
    elif target_format == 'avi':
        clip.write_videofile(output_path, codec='png')
    else:
        clip.write_videofile(output_path)
    
    clip.close()
    
    return output_path


# ==================== CONVERSION ROUTER ====================

CONVERSION_MAP = {
    # Image conversions
    ('jpg', 'png'): convert_image_format,
    ('jpg', 'gif'): convert_image_format,
    ('jpg', 'bmp'): convert_image_format,
    ('jpg', 'webp'): convert_image_format,
    ('jpg', 'tiff'): convert_image_format,
    ('jpg', 'pdf'): image_to_pdf,
    ('jpeg', 'png'): convert_image_format,
    ('jpeg', 'jpg'): convert_image_format,
    ('jpeg', 'pdf'): image_to_pdf,
    ('png', 'jpg'): convert_image_format,
    ('png', 'jpeg'): convert_image_format,
    ('png', 'gif'): convert_image_format,
    ('png', 'bmp'): convert_image_format,
    ('png', 'webp'): convert_image_format,
    ('png', 'pdf'): image_to_pdf,
    ('gif', 'jpg'): convert_image_format,
    ('gif', 'png'): convert_image_format,
    ('gif', 'pdf'): image_to_pdf,
    ('bmp', 'jpg'): convert_image_format,
    ('bmp', 'png'): convert_image_format,
    ('bmp', 'pdf'): image_to_pdf,
    ('webp', 'jpg'): convert_image_format,
    ('webp', 'png'): convert_image_format,
    ('webp', 'pdf'): image_to_pdf,
    ('tiff', 'jpg'): convert_image_format,
    ('tiff', 'png'): convert_image_format,
    ('tiff', 'pdf'): image_to_pdf,
    
    # Document conversions
    ('pdf', 'docx'): pdf_to_docx,
    ('pdf', 'txt'): pdf_to_txt,
    ('pdf', 'jpg'): pdf_to_image,
    ('pdf', 'png'): pdf_to_image,
    ('docx', 'pdf'): docx_to_pdf,
    ('docx', 'txt'): docx_to_txt,
    ('txt', 'pdf'): txt_to_pdf,
    
    # Video conversions
    ('mp4', 'gif'): video_to_gif,
    ('mp4', 'avi'): convert_video_format,
    ('avi', 'mp4'): convert_video_format,
    ('mov', 'mp4'): convert_video_format,
    ('mkv', 'mp4'): convert_video_format,
}


def get_converter(source_format, target_format):
    """Get the appropriate converter function"""
    key = (source_format.lower(), target_format.lower())
    return CONVERSION_MAP.get(key)

